import csv
import streamlit as st
from docx import Document 
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches,Cm


# doc structuring functions
def format_heading(level :int, text :str, font_name :str, font_size :int):
    heading = st.session_state.doc.add_heading(level=level)
    heading_run = heading.add_run(text)
    heading_run.font.name = font_name
    # 'Times New Roman'
    heading_run.font.size = Pt(font_size)
    # Align the paragraph to the left
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def format_heading2(level :int, text :str, font_name :str, font_size :int):
    heading = st.session_state.doc.add_heading(level=level)
    heading_run = heading.add_run(text)
    heading_run.font.name = font_name
    # 'Times New Roman'
    heading_run.font.size = Pt(font_size)
    # Align the paragraph to the left
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER


def format_paragraph(text :str, font_name :str, font_size :int, space_after:int, bold_flag= False, align = WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph = st.session_state.doc.add_paragraph(text)
    paragraph.style.font.name = font_name
    paragraph.style.font.size = Pt(font_size)
    paragraph.alignment = align
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(space_after)
    paragraph.runs[0].bold=bold_flag


def format_image(image_path :str, width :int):
    paragraph = st.session_state.doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(6), height=Cm(8))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def format_table(csv_path :str):
    with open(csv_path, 'r') as csvfile:
        csvreader = csv.reader(csvfile)
        # Read the header row to get the number of columns
        header = next(csvreader)
        num_cols = len(header)
        # Add a table with the number of columns from the header
        table = st.session_state.doc.add_table(rows=1, cols=num_cols)
        table.style = 'Colorful List'

        # Populate the header row
        for i, header_text in enumerate(header):
            table.cell(0, i).text = header_text

        # Populate the rest of the table with data from the CSV file
        for row in csvreader:
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row):
                row_cells[i].text = cell_value 